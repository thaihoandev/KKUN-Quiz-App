import fitz  # PyMuPDF
from docx import Document
import pytesseract
from PIL import Image
import io
import re
import pdfplumber
import cv2
import numpy as np
import os

def extract_text_from_file(file_path, ocr_lang='eng+vie'):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    if not file_path.lower().endswith(('.pdf', '.docx')):
        raise ValueError(f"Unsupported file type: {file_path}")
    
    text = ''
    if file_path.lower().endswith('.pdf'):
        with pdfplumber.open(file_path) as pdf:
            doc = fitz.open(file_path)
            for page_num, (plumber_page, fitz_page) in enumerate(zip(pdf.pages, doc)):
                # Text extraction
                page_text = fitz_page.get_text()
                if page_text.strip():
                    text += page_text + '\n'
                else:
                    # OCR for scanned pages
                    pix = fitz_page.get_pixmap(dpi=300)
                    img = Image.open(io.BytesIO(pix.tobytes()))
                    page_text = pytesseract.image_to_string(img, lang=ocr_lang)
                    text += page_text + '\n'
                
                # Table extraction
                for table in plumber_page.extract_tables():
                    for row in table:
                        text += '\t'.join([str(cell) for cell in row if cell]) + '\n'
                
                # Image extraction and OCR
                images = fitz_page.get_images(full=True)
                for img_index, img in enumerate(images):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image['image']
                    image = Image.open(io.BytesIO(image_bytes))
                    # Preprocess image
                    image_np = np.array(image)
                    image_np = cv2.GaussianBlur(image_np, (5, 5), 0)  # Denoise
                    gray = cv2.cvtColor(image_np, cv2.COLOR_BGR2GRAY)
                    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
                    enhanced = clahe.apply(gray)
                    thresh = cv2.threshold(enhanced, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
                    image_text = pytesseract.image_to_string(thresh, lang=ocr_lang)
                    text += image_text + '\n'
            doc.close()
    elif file_path.lower().endswith('.docx'):
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + '\n'
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + '\t'
                text += '\n'
    
    # Clean text
    text = re.sub(r'\s+', ' ', text).strip()
    return text